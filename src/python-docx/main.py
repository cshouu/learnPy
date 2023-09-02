from docx import Document

document = Document()

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2)

document.add_page_break()

table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

document.save('test.docx')